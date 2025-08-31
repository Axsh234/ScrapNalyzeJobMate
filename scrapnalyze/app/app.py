from fastapi import FastAPI, Request, Depends, Form
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, Column, Integer, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from pathlib import Path
from docx import Document
import json
from docx import Document  # <- This is still correct once you install the right package
import os

# ---------------- Database Configuration ----------------
DB_CONFIG = {
    'host': 'localhost',
    'port': 5432,
    'database': 'scrapnalyze',
    'user': 'postgres',
    'password': '1234'
}
DATABASE_URL = f"postgresql://{DB_CONFIG['user']}:{DB_CONFIG['password']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine, autocommit=False, autoflush=False)
Base = declarative_base()

# ---------------- Models ----------------
class Job(Base):
    __tablename__ = "myjob"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(Text, nullable=True)
    salary = Column(Text, nullable=True)
    date_posted = Column(Text, nullable=True)
    location = Column(Text, nullable=True)
    closing_date = Column(Text, nullable=True)
    link = Column(Text, nullable=True)  # External application link

Base.metadata.create_all(bind=engine)  # Ensure table exists

# ---------------- FastAPI App Initialization ----------------
app = FastAPI()
BASE_DIR = Path(__file__).resolve().parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR.parent / "static")), name="static")

# ---------------- Middleware ----------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

# ---------------- Dependency ----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()




####################### ---------------- Routes ----------------##################################

@app.get("/", response_class=HTMLResponse)
def root(request: Request, db: Session = Depends(get_db)):
    job_count = db.query(Job).count()
    latest_jobs = db.query(Job).order_by(Job.id.desc()).limit(5).all()
    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "job_count": job_count,
        "latest_jobs": latest_jobs
    })

################################ Job details page   ###########################################################################

from fastapi import HTTPException

@app.get("/jobs/{job_id}", response_class=HTMLResponse)
def job_detail(request: Request, job_id: int, db: Session = Depends(get_db)):
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    return templates.TemplateResponse("job_detail.html", {
        "request": request,
        "job": job
    })



########################################################################################


from fastapi.responses import JSONResponse

@app.get("/jobs/autocomplete")
def autocomplete_jobs(q: str = "", db: Session = Depends(get_db)):
    if not q:
        return JSONResponse(content=[])
    # Search job titles that contain 'q' (case-insensitive), limit 10 results
    results = (
        db.query(Job.title)
        .filter(Job.title.ilike(f"%{q}%"))
        .limit(10)
        .all()
    )
    suggestions = [r[0] for r in results]  # extract titles from tuples
    return JSONResponse(content=suggestions)


from sqlalchemy import asc, desc

@app.get("/jobs", response_class=HTMLResponse)
def list_jobs(
    request: Request,
    q: str = "",
    location: str = "",
    sort_by: str = "date_posted",  # default sort by date_posted
    order: str = "desc",  # ascending or descending
    db: Session = Depends(get_db),
):
    query = db.query(Job)
    if q:
        query = query.filter(Job.title.ilike(f"%{q}%"))
    if location:
        query = query.filter(Job.location.ilike(f"%{location}%"))

    # Mapping sort_by to actual Job columns
    sort_options = {
        "title": Job.title,
        "date_posted": Job.date_posted,
        "salary": Job.salary,
        "location": Job.location,
    }

    sort_column = sort_options.get(sort_by, Job.date_posted)

    if order == "asc":
        query = query.order_by(asc(sort_column))
    else:
        query = query.order_by(desc(sort_column))

    jobs = query.all()

    return templates.TemplateResponse("jobs.html", {
        "request": request,
        "jobs": jobs,
        "q": q,
        "location": location,
        "sort_by": sort_by,
        "order": order,
    })


@app.get("/jobs/autocomplete/location")
def autocomplete_locations(q: str = "", db: Session = Depends(get_db)):
    if not q:
        return JSONResponse(content=[])
    results = (
        db.query(Job.location)
        .filter(Job.location.ilike(f"%{q}%"))
        .distinct()
        .limit(10)
        .all()
    )
    suggestions = [r[0] for r in results]
    return JSONResponse(content=suggestions)


###################################

@app.get("/about", response_class=HTMLResponse)
def about_page(request: Request, db: Session = Depends(get_db)):
    total_jobs = db.query(Job).count()
    unique_locations = db.query(Job.location).distinct().count()
    return templates.TemplateResponse("about.html", {
        "request": request,
        "total_jobs": total_jobs,
        "unique_locations": unique_locations
    })


#######################################################################


@app.get("/career-advice", response_class=HTMLResponse)
def career_advice(request: Request):
    advice_file = Path("data/career_tips.json")
    tips = json.loads(advice_file.read_text(encoding="utf-8")) if advice_file.exists() else []
    return templates.TemplateResponse("career_advice.html", {
        "request": request,
        "tips": tips
    })
#################################################################################


@app.post("/cv-generator")
def generate_cv(name: str = Form(...), skills: str = Form(...)):
    doc = Document()
    doc.add_heading(f"Curriculum Vitae - {name}", level=1)
    doc.add_paragraph("Skills:")
    for skill in skills.split(','):
        doc.add_paragraph(f"• {skill.strip()}")
    
    # Ensure directory exists
    out_dir = Path("generated_cvs")
    out_dir.mkdir(parents=True, exist_ok=True)
    filepath = out_dir / f"{name.replace(' ', '_')}_cv.docx"
    doc.save(filepath)
    
    return FileResponse(path=filepath, filename=f"{name}_cv.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


#############################################################################################


@app.post("/cv-job-matcher", response_class=HTMLResponse)
async def cv_job_matcher(request: Request, skills: str = Form(...), db: Session = Depends(get_db)):
    skill_set = [s.strip().lower() for s in skills.split(",")]
    matched_jobs = []

    for job in db.query(Job).all():
        for skill in skill_set:
            if skill in job.title.lower():
                matched_jobs.append(job)
                break

    return templates.TemplateResponse("cv_matcher.html", {
        "request": request,
        "matched_jobs": matched_jobs,
        "skills": skill_set
    })


# uvicorn app:app --reload
